# Модуль содержит утилиты для работы с PDF и DOCX: извлечение текста, метаданных, валидация файлов, поддержка форматов.
import fitz  # PyMuPDF
from pathlib import Path
from typing import Optional, Dict, Any
import hashlib
import os
from loguru import logger
from docx import Document as DocxDocument

logger = logger.bind(name=__name__)


class PDFProcessingError(Exception):
    """Исключение для ошибок обработки PDF"""
    pass


class DocxProcessingError(Exception):
    """Исключение для ошибок обработки DOCX"""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Извлекает текст из PDF файла
    
    Args:
        file_path: Путь к PDF файлу
        
    Returns:
        Извлеченный текст
        
    Raises:
        PDFProcessingError: Если не удалось обработать PDF
    """
    try:
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        
        doc.close()
        return text
        
    except Exception as e:
        raise PDFProcessingError(f"Ошибка извлечения текста из PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Извлекает текст из DOCX файла
    
    Args:
        file_path: Путь к DOCX файлу
        
    Returns:
        Извлеченный текст
        
    Raises:
        DocxProcessingError: Если не удалось обработать DOCX
    """
    try:
        doc = DocxDocument(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
        
    except Exception as e:
        raise DocxProcessingError(f"Ошибка извлечения текста из DOCX: {e}")


def extract_text_from_file(file_path: str) -> str:
    """
    Извлекает текст из файла (PDF или DOCX)
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        Извлеченный текст
    """
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        return extract_text_from_pdf(str(file_path))
    elif file_path.suffix.lower() == '.docx':
        return extract_text_from_docx(str(file_path))
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {file_path.suffix}")


def get_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """
    Получает метаданные PDF файла
    
    Args:
        file_path: Путь к PDF файлу
        
    Returns:
        Словарь с метаданными
    """
    try:
        doc = fitz.open(file_path)
        metadata = doc.metadata
        doc.close()
        
        # Добавляем информацию о файле
        file_info = Path(file_path)
        metadata.update({
            "file_name": file_info.name,
            "file_size": file_info.stat().st_size,
            "file_extension": file_info.suffix.lower(),
            "file_hash": calculate_file_hash(file_path)
        })
        
        return metadata
        
    except Exception as e:
        logger.warning("Не удалось получить метаданные PDF", file=file_path, error=str(e))
        return {
            "file_name": Path(file_path).name,
            "file_extension": Path(file_path).suffix.lower(),
            "error": str(e)
        }


def get_docx_metadata(file_path: str) -> Dict[str, Any]:
    """
    Получает метаданные DOCX файла
    
    Args:
        file_path: Путь к DOCX файлу
        
    Returns:
        Словарь с метаданными
    """
    try:
        doc = DocxDocument(file_path)
        file_info = Path(file_path)
        
        metadata = {
            "file_name": file_info.name,
            "file_size": file_info.stat().st_size,
            "file_extension": file_info.suffix.lower(),
            "file_hash": calculate_file_hash(file_path),
            "paragraphs_count": len(doc.paragraphs),
            "sections_count": len(doc.sections)
        }
        
        # Пытаемся получить свойства документа
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.created:
                metadata["created_date"] = str(props.created)
            if props.modified:
                metadata["modified_date"] = str(props.modified)
        
        return metadata
        
    except Exception as e:
        logger.warning("Не удалось получить метаданные DOCX", file=file_path, error=str(e))
        return {
            "file_name": Path(file_path).name,
            "file_extension": Path(file_path).suffix.lower(),
            "error": str(e)
        }


def get_file_metadata(file_path: str) -> Dict[str, Any]:
    """
    Получает метаданные файла (PDF или DOCX)
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        Словарь с метаданными
    """
    file_path = Path(file_path)
    
    if file_path.suffix.lower() == '.pdf':
        return get_pdf_metadata(str(file_path))
    elif file_path.suffix.lower() == '.docx':
        return get_docx_metadata(str(file_path))
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {file_path.suffix}")


def calculate_file_hash(file_path: str) -> str:
    """
    Вычисляет хеш файла
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        MD5 хеш файла
    """
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def is_supported_file(file_path: str) -> bool:
    """
    Проверяет, поддерживается ли тип файла
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        True если файл поддерживается
    """
    supported_extensions = {'.pdf', '.docx'}
    return Path(file_path).suffix.lower() in supported_extensions


def validate_pdf_file(file_path: str) -> bool:
    """
    Проверяет, является ли файл валидным PDF
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        True если файл валидный PDF
    """
    try:
        file_path = Path(file_path)
        
        if not file_path.exists():
            return False
        
        if not file_path.suffix.lower() == '.pdf':
            return False
        
        # Пытаемся открыть файл как PDF
        with fitz.open(file_path) as doc:
            if len(doc) == 0:
                return False
        
        return True
        
    except Exception:
        return False 